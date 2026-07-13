from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os


def export_docx(topic: str, report: str) -> str:
    os.makedirs("exports", exist_ok=True)

    safe_topic = "".join(
        c if c.isalnum() or c in (" ", "_", "-") else ""
        for c in topic
    ).strip()

    filename = f"{safe_topic}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join("exports", filename)

    doc = Document()

    # -------------------------
    # Cover Page
    # -------------------------

    title = doc.add_heading("AI Multi-Agent Research Assistant", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Professional Research Report")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()

    topic_para = doc.add_paragraph()
    topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = topic_para.add_run("Research Topic\n")
    run.bold = True
    run.font.size = Pt(16)

    run = topic_para.add_run(topic)
    run.font.size = Pt(20)

    doc.add_paragraph()

    generated = doc.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.add_run(
        f"Generated on:\n{datetime.now().strftime('%d %B %Y, %I:%M %p')}"
    )

    doc.add_paragraph()
    doc.add_paragraph("_" * 70)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(
        "Generated using LangGraph • Ollama • DuckDuckGo Search • Qwen3"
    ).italic = True

    # New Page
    doc.add_page_break()

    # -------------------------
    # Report Title
    # -------------------------

    heading = doc.add_heading(topic, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # -------------------------
    # Report Body
    # -------------------------

    for line in report.split("\n"):

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)

        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)

        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)

        elif line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(
                line[2:],
                style="List Bullet"
            )

        elif line.strip() == "":
            doc.add_paragraph()

        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)

    doc.save(filepath)

    return filepath
